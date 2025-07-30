import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter


def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""
    loader = None
    if input_type == "Link":
        # Handle multiple URLs
        if isinstance(input_data, list):
            # Filter out empty URLs
            valid_urls = [url for url in input_data if url and url.strip()]
            if not valid_urls:
                raise ValueError("No valid URLs provided")
            loader = WebBaseLoader(valid_urls)
        else:
            loader = WebBaseLoader([input_data])
        documents = loader.load()
    elif input_type == "PDF":
        if hasattr(input_data, 'read'):  # File-like object
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        elif isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if hasattr(input_data, 'read'):  # File-like object
            doc = Document(BytesIO(input_data.read()))
        elif isinstance(input_data, BytesIO):
            doc = Document(input_data)
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if hasattr(input_data, 'read'):  # File-like object
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    # Use smaller chunk size to handle context length
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)
    return texts
