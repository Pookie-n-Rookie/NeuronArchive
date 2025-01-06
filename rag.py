import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint

from secret_api_keys import huggingface_api_key
os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key

def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    # Use smaller chunk size to handle context length
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}
    
    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)
    return vector_store

def answer_question(vectorstore, query):
    """Answers a question based on the provided vectorstore."""
    llm = HuggingFaceEndpoint(
        repo_id='meta-llama/Meta-Llama-3-8B-Instruct',
        token=huggingface_api_key,
        temperature=0.8,
        max_new_tokens=256  
    )
    
   
    qa = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={'k': 3})  
    )

    answer = qa({"query": query})
    return answer

def main():
    st.set_page_config(page_title="Document Q&A Assistant", layout="wide")
    
    st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        margin-top: 1rem;
    }
    .css-1d391kg {
        padding: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)

    st.title("📚 Document Q&A Assistant")
    st.markdown("---")

    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("Document Input")
        input_type = st.selectbox("Select Input Type", 
                                ["Text", "PDF", "DOCX", "TXT", "Link"],
                                help="Choose the type of document you want to analyze")
        
        if input_type == "Link":
            number_input = st.number_input("Number of Links", 
                                         min_value=1, max_value=20, 
                                         step=1,
                                         help="Enter how many URLs you want to analyze")
            input_data = []
            for i in range(number_input):
                url = st.text_input(f"URL {i+1}", 
                                  help="Enter the complete URL including http:// or https://")
                input_data.append(url)
        elif input_type == "Text":
            input_data = st.text_area("Enter your text", 
                                    height=200,
                                    help="Paste or type your text here")
        else:
            input_data = st.file_uploader(f"Upload {input_type} file", 
                                        type=[input_type.lower()],
                                        help=f"Upload your {input_type} document")

        if st.button("Process Document", type="primary"):
            try:
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
                st.success("Document processed successfully!")
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")

    with col2:
        st.subheader("Ask Questions")
        if "vectorstore" in st.session_state:
            query = st.text_input("What would you like to know about your document?",
                                help="Ask any question about the content of your document")
            
            if st.button("Get Answer", type="primary"):
                try:
                    answer = answer_question(st.session_state["vectorstore"], query)
                    st.markdown("### Answer")
                    st.markdown(answer["result"])
                except Exception as e:
                    st.error(f"Error generating answer: {str(e)}")
        else:
            st.info("Please process a document first before asking questions.")

        if "answers" in st.session_state:
            st.markdown("### Previous Questions & Answers")
            for q, a in st.session_state["answers"]:
                with st.expander(q):
                    st.write(a)

if __name__ == "__main__":
    main()